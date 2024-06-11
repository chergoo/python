from configparser import ConfigParser
import glob
import os
from tkinter import Tk
from tkinter import filedialog
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt
from docx.oxml.ns import qn

# 自动查找并读取 ini 文件
def find_ini_file(directory):
    ini_files = glob.glob(os.path.join(directory, '*.ini'))
    if ini_files:
        return ini_files[0]  # 返回第一个找到的 ini 文件
    else:
        raise FileNotFoundError("No ini file found in the specified directory.")

# 弹窗选择文件夹路径
def select_directory():
    root = Tk()
    root.withdraw()  # 隐藏主窗口
    directory = filedialog.askdirectory()  # 弹出选择文件夹对话框
    root.destroy()  # 销毁主窗口
    if not directory:
        raise Exception("No directory selected.")
    return directory

def read_ini_file(file_path):
    with open(file_path, 'r') as c_file:
        config = c_file.read()
    return config

# 查找并读取 dat 文件
def find_dat_files(directory):
    search_pattern = os.path.join(directory, f"*.dat")
    return glob.glob(search_pattern)

def read_dat_file(file_path):
    with open(file_path, 'r') as file:
        data = file.read()
        # print("dat文件已读取")
    return data  

#提取部分文字
def extract_text(dat_contents, start_marker, end_marker):
    sections = dat_contents.split('\n')  # 假设段落之间以双换行分隔
    extracted_sections = []
    capturing = False
    
    for section in sections:
        if start_marker in section:
            capturing = True
        if capturing:
            extracted_sections.append(section)
        if end_marker in section:
            capturing = False
            break
    return '\n'.join(extracted_sections).strip()


# 主程序
def main():
    # 选择文件夹路径
    directory = select_directory()
    
    # 自动查找并读取 ini 文件
    ini_file_path = find_ini_file(directory)
    ini_file = read_ini_file(ini_file_path)
    # print("ini",ini_file)

    # 查找并读取 dat 文件
    dat_files = find_dat_files(directory)
    dat_contents = [read_dat_file(file) for file in dat_files]
    # print(dat_contents)
        
    # document = create_word_doc(extracted_text, columns =2)
    # 创建一个新的Word文档对象
    doc = Document()

    section = doc.add_section(WD_SECTION.CONTINUOUS)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set('num', str(2))

    # 读取并写入INI文件内容
    ini_content = read_ini_file(ini_file_path)
    doc.add_heading('INI File Content', level=1)
    doc.add_paragraph(ini_content)

    

    # 读取并写入DAT文件内容
    for dat_file in dat_files:
        dat_file_path = os.path.join(directory, dat_file)
        start_marker = 'DATA'
        end_marker = '[END] of [DATA]'
        dat_contents = read_dat_file(dat_file_path)
        extracted_text = extract_text(dat_contents, start_marker, end_marker)
        doc.add_heading(f'DAT File: {dat_file}', level=1)
        doc.add_paragraph(extracted_text)

    # 保存Word文档
    output_docx_path = 'C:\\Users\\gchen\\Desktop\\output.docx'
    doc.save(output_docx_path)
    print("文件已输出至",output_docx_path)

if __name__ == "__main__":
    main()
